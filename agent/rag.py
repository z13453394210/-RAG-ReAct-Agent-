"""
rag-knowledge-agent — RAG 引擎

处理文档接入（PDF, DOCX, TXT, MD）、文本分块、
使用本地 HuggingFace 模型做向量化、基于 numpy 的本地向量存储。
"""

import os
import json
import pickle
import hashlib
from pathlib import Path
from typing import List, Optional

import torch
import torch.nn.functional as F
import numpy as np
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyMuPDFLoader, TextLoader
from langchain_core.documents import Document
from transformers import AutoModel, AutoTokenizer
from dotenv import load_dotenv

# 从项目根目录加载 .env
_env_path = Path(__file__).resolve().parent.parent / ".env"
if _env_path.exists():
    load_dotenv(dotenv_path=str(_env_path))

# 默认使用 HuggingFace 中国镜像
os.environ.setdefault("HF_ENDPOINT", "https://hf-mirror.com")

# ---------------------------------------------------------------------------
# 路径
# ---------------------------------------------------------------------------
VECTOR_STORE_DIR = Path(__file__).resolve().parent.parent / "vector_store" / "chroma_db"
UPLOAD_DIR = Path(__file__).resolve().parent.parent / "docs" / "uploads"

VECTOR_STORE_DIR.mkdir(parents=True, exist_ok=True)
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

DATA_FILE = VECTOR_STORE_DIR / "vectors.pkl"


# ---------------------------------------------------------------------------
# 本地嵌入模型（基于 transformers + PyTorch，无需 API Key）
# ---------------------------------------------------------------------------
_EMBEDDING_TOKENIZER = None
_EMBEDDING_MODEL = None


def _load_embedding_model():
    global _EMBEDDING_TOKENIZER, _EMBEDDING_MODEL
    if _EMBEDDING_MODEL is None:
        model_name = os.getenv("EMBEDDING_MODEL", "BAAI/bge-small-zh-v1.5")
        _EMBEDDING_TOKENIZER = AutoTokenizer.from_pretrained(model_name, local_files_only=True)
        _EMBEDDING_MODEL = AutoModel.from_pretrained(model_name, local_files_only=True)
        _EMBEDDING_MODEL.eval()


def _mean_pooling(token_embeddings, attention_mask):
    mask = attention_mask.unsqueeze(-1).expand(token_embeddings.size()).float()
    pooled = torch.sum(token_embeddings * mask, 1) / torch.clamp(mask.sum(1), min=1e-9)
    return pooled


class LocalEmbeddings:
    def __init__(self):
        _load_embedding_model()

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        encoded = _EMBEDDING_TOKENIZER(texts, padding=True, truncation=True, return_tensors="pt", max_length=512)
        with torch.no_grad():
            output = _EMBEDDING_MODEL(**encoded)
        pooled = _mean_pooling(output[0], encoded["attention_mask"])
        pooled = F.normalize(pooled, p=2, dim=1)
        return pooled.numpy().tolist()

    def embed_query(self, text: str) -> List[float]:
        return self.embed_documents([text])[0]


_embeddings_instance = None


def get_embeddings():
    global _embeddings_instance
    if _embeddings_instance is None:
        _embeddings_instance = LocalEmbeddings()
    return _embeddings_instance


# ---------------------------------------------------------------------------
# 基于 numpy 的本地向量存储（替代 ChromaDB，无 C 扩展依赖）
# ---------------------------------------------------------------------------
class SimpleVectorStore:
    """使用 numpy 的轻量向量存储。持久化为 pickle 文件。"""

    def __init__(self):
        self.documents: List[dict] = []
        self._load()

    def _load(self):
        if DATA_FILE.exists():
            try:
                with open(DATA_FILE, "rb") as f:
                    self.documents = pickle.load(f)
            except Exception:
                self.documents = []

    def _save(self):
        with open(DATA_FILE, "wb") as f:
            pickle.dump(self.documents, f)

    def add_documents(self, chunks: List[Document], embeddings: List[List[float]]):
        for chunk, emb in zip(chunks, embeddings):
            self.documents.append({
                "text": chunk.page_content,
                "metadata": chunk.metadata,
                "embedding": emb,
            })
        self._save()

    def similarity_search(self, query_emb: List[float], k: int = 5) -> List[Document]:
        if not self.documents:
            return []
        q = np.array(query_emb)
        docs_arr = np.array([d["embedding"] for d in self.documents])
        scores = np.dot(docs_arr, q)
        top_k = min(k, len(scores))
        indices = np.argsort(scores)[::-1][:top_k]
        results = []
        for i in indices:
            d = self.documents[i]
            results.append(Document(page_content=d["text"], metadata=d["metadata"]))
        return results

    def get_all_sources(self) -> List[str]:
        sources = set()
        for d in self.documents:
            src = d["metadata"].get("source", "")
            if src:
                sources.add(Path(src).name)
        return sorted(sources)

    def clear(self):
        self.documents = []
        if DATA_FILE.exists():
            DATA_FILE.unlink()


_store_instance = None


def get_store() -> SimpleVectorStore:
    global _store_instance
    if _store_instance is None:
        _store_instance = SimpleVectorStore()
    return _store_instance


# ---------------------------------------------------------------------------
# 文档加载
# ---------------------------------------------------------------------------
def _load_pdf(path: str) -> List[Document]:
    return PyMuPDFLoader(path).load()


def _load_docx(path: str) -> List[Document]:
    from docx import Document as DocxDocument
    doc = DocxDocument(path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return [Document(page_content=text, metadata={"source": path})]


def _load_text(path: str) -> List[Document]:
    return TextLoader(path, encoding="utf-8").load()


def _load_markdown(path: str) -> List[Document]:
    with open(path, "r", encoding="utf-8") as f:
        text = f.read()
    return [Document(page_content=text, metadata={"source": path})]


LOADER_MAP = {
    ".pdf": _load_pdf,
    ".docx": _load_docx,
    ".txt": _load_text,
    ".md": _load_markdown,
}


def load_document(file_path: str) -> List[Document]:
    ext = Path(file_path).suffix.lower()
    loader = LOADER_MAP.get(ext)
    if loader is None:
        raise ValueError(f"不支持的文件格式: {ext}")
    docs = loader(file_path)
    for d in docs:
        d.metadata.setdefault("source", file_path)
    return docs


# ---------------------------------------------------------------------------
# 分块
# ---------------------------------------------------------------------------
def chunk_documents(docs: List[Document], chunk_size: int = 1000, chunk_overlap: int = 200) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""],
    )
    return splitter.split_documents(docs)


# ---------------------------------------------------------------------------
# 接入
# ---------------------------------------------------------------------------
def ingest_document(file_path: str) -> int:
    raw_docs = load_document(file_path)
    chunks = chunk_documents(raw_docs)

    emb = get_embeddings()
    texts = [c.page_content for c in chunks]
    vectors = emb.embed_documents(texts)

    store = get_store()
    store.add_documents(chunks, vectors)
    return len(chunks)


def list_indexed_sources() -> List[str]:
    store = get_store()
    return store.get_all_sources()


def delete_all_documents() -> int:
    store = get_store()
    count = len(store.documents)
    store.clear()
    return count


# ---------------------------------------------------------------------------
# 检索
# ---------------------------------------------------------------------------
def search_knowledge_base(query: str, k: int = 5) -> List[Document]:
    emb = get_embeddings()
    query_vec = emb.embed_query(query)
    store = get_store()
    return store.similarity_search(query_vec, k=k)


def format_retrieved_docs(docs: List[Document]) -> str:
    lines: List[str] = []
    for i, doc in enumerate(docs, 1):
        source = Path(doc.metadata.get("source", "unknown")).name
        lines.append(f"[来源 {i}] ({source})\n{doc.page_content.strip()}\n")
    return "\n---\n".join(lines)